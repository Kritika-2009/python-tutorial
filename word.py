pip install python-docx


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Word document
doc = Document()

# Title
title = doc.add_heading('The Bright Future Times', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('The Official Newsletter of Sunrise International School')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Edition info
edition_info = doc.add_paragraph('Edition: April 2025 | Volume: 12 | Issue: 1\nLocation: Gurugram, Haryana | www.sunriseinternationalschool.edu.in')
edition_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Principal's Message
doc.add_heading('A Message from the Principal', level=1)
principal_message = (
    "Dear Students, Parents, and Faculty,\n\n"
    "As we step into a vibrant new academic session, I extend my warmest greetings to all of you. "
    "The past year has been filled with learning, growth, and resilience. I am proud of our students’ achievements "
    "and the unwavering support of our teachers and parents.\n\n"
    "This edition of The Bright Future Times celebrates our school’s journey so far—covering academic successes, "
    "co-curricular highlights, and the amazing creativity of our students. Let's continue to foster an environment of "
    "curiosity, compassion, and courage.\n\n"
    "Here’s to a year of excellence and new opportunities!\n\n"
    "Warm regards,\nMrs. Anjali Mehra\nPrincipal, Sunrise International School"
)
doc.add_paragraph(principal_message)

# Highlights Section
doc.add_heading('In This Edition:', level=1)
highlights = [
    "Highlights from the Annual Day & Science Exhibition",
    "Meet the Top Achievers: Board Exam Results 2024",
    "Inter-House Sports Competition: Victory & Team Spirit",
    "Student Corner: Poems, Art, and Short Stories",
    "Teacher Spotlight: Interviews & New Faculty",
    "Upcoming Events: Excursions, Workshops, and More!"
]
for item in highlights:
    doc.add_paragraph(f"✅ {item}", style='List Bullet')

# Editor's Note
doc.add_heading("Editor’s Note:", level=1)
editor_note = (
    "Welcome to the April 2025 edition! We’ve captured the spirit of our school in every story and picture shared. "
    "A big thank you to the student editorial team and all contributors. Keep reading, keep creating!\n\n"
    "— The Editorial Team"
)
doc.add_paragraph(editor_note)

# Save the document
doc.save("School_Newsletter_Front_Page.docx")
